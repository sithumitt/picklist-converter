import io
import re
import pdfplumber
import pandas as pd
from docx import Document
import streamlit as st

# Product Mapping Dictionary
PRODUCT_MAPPING = {
    "aloe vera drink 1l": "කෝමාරිකා බීම ලීටර් 1",
    "aloe vera drink 200ml": "කෝමාරිකා බීම 200",
    "aloe vera drink 500ml": "කෝමාරිකා බීම 500",
    "assortment 150gm": "ඇසෝඩ්මන්ට් බිස්කට් 150",
    "assortment 300gm": "ඇසෝඩ්මන්ට් බිස්කට් 300",
    "butter cookies 120gm": "බටර් කුකීස් 120",
    "cheese cuts tin 210gm": "චීස් බිස්කට් රතු ටින්",
    "chilli paste 60g": "චිලී පේස්ට් 60",
    "choco shorties 220gm": "චොකෝ ශෝටීස් 220",
    "chocolate chips cookies 120gm": "චොකලට් චිප් කුකීස් 120",
    "chocolate chips cookies 65gm": "චොකලට් චිප් කුකීස් 65",
    "chocolate cream 500gm": "චොකලට් ක්‍රීම්",
    "chocolate wafer 40gm": "චොකලට වේපස් 40",
    "chocolate wafer 90gm": "චොකලට වේපස් 90",
    "chocolate wafer 200gm": "චොකලට වේපස් 200",
    "chocolate marie 100gm":"චොකලට් මාරි 100GM",
    "chocolate wafer 360gm x 3pks": "චොකලට වේපස් 360",
    "cookie assortment blue 330gm": "කුකි ඇසෝඩ්මන්ට් නිල් 330",
    "coffee wafer 200gm": "කෝපි වේපර්ස්",
    "creamy choc 90gm": "ක්‍රීමි චොක් 90",
    "creamy choc 210gm": "ක්‍රීමි චොක් 210",
    "creamy choc 360gm x 4pks": "ක්‍රීමි චොක් 360",
    "creamy vanilla 360gm": "ක්‍රීම් වැනිලා 360",
    "falooda wafer 200gm": "ෆාලූඩා වේපස් 200",
    "ginger 80gm": "ඉඟුරු බිස්කට් 80",
    "ginger 250gm": "ඉඟුරු බිස්කට් 250",
    "green apple sparkling 250ml": "ඇපල් ස්පාක්ලින්",
    "kist cream cracker 125gm": "ක්‍රීම් ක්‍රැකර් 125",
    "kist cream cracker 500gm": "ක්‍රීම් ක්‍රැකර් 500",
    "kist gift assortment 400gm": "ගීෆ්ට් ඇසෝඩ්මන්ට් 400",
    "kithul treacle 340ml": "කිතුල් පැණි 340",
    "knuckles water 500ml": "වතුර 500",
    "knuckles water 1000ml": "වතුර 1000",
    "knuckles water 1500ml": "වතුර 1500",
    "lemon & mint nectar 1l": "ලෙමන් සහ මින්ට් බීම 1l",
    "lemon puff 100gm": "ලෙමන් පෆ් 100",
    "lemon puff 200gm": "ලෙමන් පෆ් 200",
    "magic choco fun 40gm": "චෝෆන් කහ පාට 40",
    "magic choco fun 100gm": "චෝෆන් කහ පාට 100",
    "magic choco nut 160gm": "චොකෝ නට් පොල් 160",
    "magic chocooh 100gm": "චොකො ඕ [සුදු පාට ]",
    "magic chocoblok chocolat 160gm": "චොකෝ බ්ලොක් චොකලට් 160",
    "magic chocoblok coffee 160gm": "චොකෝ බ්ලොක් කෝපි",
    "magic chocoblok vanilla 160gm": "චොකෝ බ්ලොක් වැනිලා",
    "magic chokstik choco 10gm": "චොකලට් චොක්ස්ටික්",
    "magic chokstik strawberry 10gm": "ස්ටෝබරි චොක්ස්ටික්",
    "magic chokstik vanilla 10gm": "වැනිලා චොක්ස්ටික්",
    "magic chonkz 120gm": "චෝන්ස් ලා නිල් 120",
    "magic chonkz 240gm": "චෝන්ස් ලා නිල් 240",
    "magic fingers 27gm": "ෆින්ගස් 27",
    "magic fingers 100gm": "ෆින්ගස් 100",
    "magic olo assortment 180gm": "ඔලෝ ඇසෝර්ට්මන් 180",
    "magic olo butter scotch 140gm": "ඕලෝ බටර් ස්කොච් 140",
    "magic olo chocolate 60gm": "ඕලෝ චොකලට් 60",
    "magic olo chocolate 140gm": "ඕලෝ චොකලට් 140",
    "magic olo classic 60gm": "ඔලෝ ක්ලැසික් 60",
    "magic olo classic 140gm": "ඔලෝ ක්ලැසික් 140",
    "magic olo strawberry 60gm": "ඔලෝ ස්ට්‍රෝබෙරි 60",
    "magic olo strawberry 140gm": "ඔලෝ ස්ට්‍රෝබෙරි 140",
    "magic olo white 140gm": "ඔලෝ වයිට් 140",
    "magic treats 90gm": "ටීට්ස් දම් පාට",
    "mango nectar 1l": "අඹ බීම 1l",
    "mango nectar 200ml": "අඹ බීම 200",
    "mango nectar 500ml": "අඹ බීම 500",
    "marie 50gm": "මාරි 50",
    "marie 100gm": "මාරි 100",
    "marie 200gm": "මාරි 200",
    "marie 500gm": "මාරි 500",
    "mayonnaise 200g pouch": "මයෝනාඊස් 200",
    "milk shorties 220gm": "කිරි ෂෝටීස් 220",
    "milki cookies 120gm": "මීල්කි කුකීස්",
    "mixed fruit jam cup 100g": "මිශ්‍ර ජෑම් කප් 100",
    "mixed fruit jam 200g": "මිශ්‍ර ජෑම් 200",
    "mixed fruit jam 300g": "මිශ්‍ර ජෑම් 300",
    "mixed fruit jam 510g": "මිශ්‍ර ජෑම් 510",
    "mixed fruit nectar 1l": "මිශ්‍ර පලතුරු බීම 1L",
    "mixed fruit nectar 200ml": "මිශ්‍ර පලතුරු බීම 200ML",
    "mixed fruit nectar 500ml": "මිශ්‍ර පලතුරු බීම 500ML",
    "nice 430gm": "නායිස් 430",
    "onion byte 30gm": "අනියන් බයිට් 30",
    "orange nectar 500ml": "දොඩම් බීම 500",
    "orange sparkling 250ml": "දොඩම් ස්පාක්ලින්",
    "passion fruit nectar 1l": "පැෂන්ෆෲට් නෙක්ටා 1L",
    "ride classic drink 250ml": "රයිට් නිල්",
    "ride redberry drink 250ml": "රයිට් රතු",
    "ride sugar free drink 250ml": "රයිට් සීනි නැති",
    "s/berry melon jam cup100g": "ස්ටෝබරි ජෑම් 100 C",
    "berry flv melon jam200g": "ස්ටෝබරි ජෑම් 200",
    "berry flv melon jam300g": "ස්ටෝබරි ජෑම් 300",
    "strawberry 200g": "ස්ටෝබරි ජෑම් 200", 
    "strawberry 300g": "ස්ටෝබරි ජෑම් 300", 
    "sesame cookies 120gm": "සෙසමිකුකීස්",
    "strawberry sparkling 250ml": "ස්ටෝබරි ස්පාක්ලින්",
    "strawberry wafer 40gm": "ස්ටෝබරි වේපස් 40",
    "strawberry wafer 90gm": "ස්ටෝබරි වේපස් 90",
    "strawberry wafer 200gm": "ස්ටෝබරි වේපස් 200",
    "strawberry wafer 360gm x 3pks": "ස්ටෝබරි වේපස් 360",
    "soya sauce squeeasy 180ml": "සෝස් පැකට් 180ml",
    "tomato sachet 15g": "සෝස් පැකට් 15",
    "tomato sauce 110gr - pouch": "සෝස් පැකට් 110",
    "tomato sauce 400g - pouch": "සෝස් පැකට් 400",
    "tomato sauce 400g ": "සෝස් වීදුරු බෝතලේ 400 ",
    "tomato sauce 200gr ": "සෝස් වීදුරු බෝතලේ 200 ",
    "vanilla wafer 40gm": "වැනිලා වේපස් 40",
    "vanilla wafer 90gm": "වැනිලා වේපස් 90",
    "vanilla wafer 200gm": "වැනිලා වේපස් 200",
    "vanilla wafer 360gm x 3pks": "වැනිලා වේපස් 360",
    "woodapple nectar 200ml": "දිවුල් බීම 200",
    "woodapple nectar 500ml": "දිවුල් බීම 500"
}

def clean_text_for_matching(text):
    if not text:
        return ""
    cleaned = text.replace('"', '').replace('/', ' ').replace('.', ' ').replace('-', ' ')
    return " ".join(cleaned.lower().split())

# App Interface Titles
st.title("📋 පික් ලිස්ට් එකේ බඩු පරිවර්තකය")
st.write("ඔබේ Picklist PDF එක සිංහල Word ගොනුවක් බවට ක්ෂණිකව පරිවර්තනය කරන්න")

# Main File Input
uploaded_file = st.file_uploader("පරිවර්තනය සඳහා PDF ගොනුවක් තෝරන්න (Select PDF File)", type=["pdf"])

if uploaded_file is not None:
    with st.spinner("දත්ත විශ්ලේෂණය කරමින් පවතී..."):
        # Setup Output Word File Structure
        doc = Document()
        doc.add_heading('පික් ලිස්ට් එකේ බඩු', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'භාණ්ඩය'
        hdr_cells[1].text = 'කේස්'
        hdr_cells[2].text = 'කෑලි'
        
        matched_count = 0
        preview_data = []
        processed_lines = set()
        
        # Parse PDF Document content
        with pdfplumber.open(uploaded_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text_content = page.extract_text()
                if not text_content:
                    continue
                
                lines = text_content.split('\n')
                for line_idx, line in enumerate(lines):
                    line_key = f"{page_num}-{line_idx}"
                    normalized_line = line.replace('"', '').strip()
                    matchable_line = clean_text_for_matching(normalized_line)
                    
                    for english_key, sinhala_val in PRODUCT_MAPPING.items():
                        matchable_key = clean_text_for_matching(english_key)
                        
                        is_match = False
                        if matchable_key in matchable_line:
                            is_match = True
                        elif "strawberry" in matchable_key and "300" in matchable_key:
                            if "strawberry" in matchable_line and "300" in matchable_line:
                                is_match = True
                                
                        if is_match and line_key not in processed_lines:
                            processed_lines.add(line_key)
                            
                            batch_match = re.search(r'\b([A-Z]{2}\d)\b', normalized_line)
                            qty1, qty2 = "0", "0"
                            
                            if batch_match:
                                pre_batch_text = normalized_line[:batch_match.start()].strip()
                                all_numbers = re.findall(r'\d+', pre_batch_text)
                                if len(all_numbers) >= 2:
                                    qty1 = all_numbers[-2]
                                    qty2 = all_numbers[-1]
                            else:
                                all_numbers = re.findall(r'\b\d+\b', normalized_line)
                                if len(all_numbers) >= 2:
                                    qty1 = all_numbers[-2]
                                    qty2 = all_numbers[-1]
                            
                            row_cells = table.add_row().cells
                            row_cells[0].text = sinhala_val
                            row_cells[1].text = qty1
                            row_cells[2].text = qty2
                            
                            preview_data.append({"භාණ්ඩය": sinhala_val, "කේස්": qty1, "කෑලි": qty2})
                            matched_count += 1
                            break 

    if matched_count > 0:
        st.success(f"🎉 සාර්ථකයි! ගැළපෙන භාණ්ඩ පේළි {matched_count} ක් සාර්ථකව පරිවර්තනය කරන ලදී.")
        
        # Render Table Preview Natively
        st.subheader("දත්ත පෙරදසුන (Data Preview)")
        df_preview = pd.DataFrame(preview_data)
        st.dataframe(df_preview, use_container_width=True)
        
        # Prepare Document download stream bytes
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        st.download_button(
            label="📥 නිපදවන ලද Word ලිපිගොනුව බාගත කරගන්න (Download Word Document)",
            data=doc_stream,
            file_name="පික්_ලිස්ට්_එකේ_බඩු.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.error("⚠️ දෝෂයකි: අප්ලෝඩ් කරන ලද PDF ගොනුවේ තිබූ කිසිදු භාණ්ඩයක් අපගේ නාමාවලිය සමඟ ගැළපුණේ නැත.")
